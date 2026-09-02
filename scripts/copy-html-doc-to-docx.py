"""Copia .doc (HTML Word) da pasta pdfs para .docx nativo do Office."""

from __future__ import annotations

import base64
import io
import re
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PDFS = ROOT / "pdfs"
TODAY = datetime.now().date()
BLUE = RGBColor(0x1E, 0x3A, 0x8A)
SLATE = RGBColor(0x1E, 0x29, 0x3B)


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = tc.makeelement(qn("w:shd"), {})
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


class HtmlDocParser(HTMLParser):
    def __init__(self, document: Document) -> None:
        super().__init__(convert_charrefs=True)
        self.doc = document
        self.skip_depth = 0
        self.bold = 0
        self.italic = 0
        self.code = 0
        self.heading = 0
        self.list_level = 0
        self.in_li = False
        self.blockquote = 0
        self.center = False
        self.para = None
        self.href: str | None = None
        self.table_rows: list[list[tuple[str, bool]]] | None = None
        self.row: list[tuple[str, bool]] | None = None
        self.cell_text: list[str] | None = None
        self.cell_header = False
        self.ignore_data = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attr = {k.lower(): (v or "") for k, v in attrs}
        if tag in {"style", "script", "title", "head"}:
            self.skip_depth += 1
            self.ignore_data = True
            return
        if self.skip_depth:
            return
        if tag in {"h1", "h2", "h3", "h4"}:
            self.flush_para()
            self.heading = int(tag[1])
            self.para = self.doc.add_heading("", level=min(self.heading, 4))
        elif tag == "p":
            self.flush_para()
            self.center = "center" in attr.get("style", "").lower()
            if self.table_rows is None:
                self.para = self.doc.add_paragraph()
                if self.center:
                    self.para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if self.blockquote:
                    self.para.paragraph_format.left_indent = Cm(1)
        elif tag == "blockquote":
            self.flush_para()
            self.blockquote += 1
        elif tag == "ul":
            self.flush_para()
            self.list_level += 1
        elif tag == "li":
            self.flush_para()
            self.in_li = True
            self.para = self.doc.add_paragraph(style="List Bullet")
        elif tag in {"strong", "b"}:
            self.bold += 1
        elif tag in {"em", "i"}:
            self.italic += 1
        elif tag == "code":
            self.code += 1
        elif tag == "a":
            self.href = attr.get("href") or None
        elif tag == "br":
            if self.para is not None:
                self.para.add_run().add_break()
            elif self.cell_text is not None:
                self.cell_text.append("\n")
        elif tag == "hr":
            self.flush_para()
            self.doc.add_paragraph("—" * 24)
        elif tag == "table":
            self.flush_para()
            self.table_rows = []
        elif tag == "tr":
            self.row = []
        elif tag in {"td", "th"}:
            self.cell_text = []
            self.cell_header = tag == "th"
        elif tag == "img":
            self.add_image(attr.get("src", ""))

    def handle_endtag(self, tag: str) -> None:
        if tag in {"style", "script", "title", "head"}:
            self.skip_depth = max(0, self.skip_depth - 1)
            self.ignore_data = self.skip_depth > 0
            return
        if self.skip_depth:
            return
        if tag in {"h1", "h2", "h3", "h4"}:
            self.heading = 0
            self.para = None
        elif tag == "p":
            self.flush_para()
            self.center = False
        elif tag == "blockquote":
            self.flush_para()
            self.blockquote = max(0, self.blockquote - 1)
        elif tag == "ul":
            self.flush_para()
            self.list_level = max(0, self.list_level - 1)
        elif tag == "li":
            self.flush_para()
            self.in_li = False
        elif tag in {"strong", "b"}:
            self.bold = max(0, self.bold - 1)
        elif tag in {"em", "i"}:
            self.italic = max(0, self.italic - 1)
        elif tag == "code":
            self.code = max(0, self.code - 1)
        elif tag == "a":
            self.href = None
        elif tag in {"td", "th"}:
            text = "".join(self.cell_text or []).strip()
            if self.row is not None:
                self.row.append((text, self.cell_header))
            self.cell_text = None
            self.cell_header = False
        elif tag == "tr":
            if self.table_rows is not None and self.row is not None:
                self.table_rows.append(self.row)
            self.row = None
        elif tag == "table":
            self.flush_table()

    def handle_data(self, data: str) -> None:
        if self.ignore_data or not data:
            return
        if self.cell_text is not None:
            self.cell_text.append(data)
            return
        if self.para is None:
            if not data.strip():
                return
            self.para = self.doc.add_paragraph()
        run = self.para.add_run(data)
        run.font.color.rgb = BLUE if self.heading else SLATE
        if self.bold:
            run.bold = True
        if self.italic:
            run.italic = True
        if self.code:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        if self.href:
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            run.underline = True

    def flush_para(self) -> None:
        self.para = None

    def add_image(self, src: str) -> None:
        match = re.match(r"data:image/([a-zA-Z0-9+]+);base64,(.+)", src, re.S)
        if not match:
            return
        raw = re.sub(r"\s+", "", match.group(2))
        try:
            blob = base64.b64decode(raw)
        except Exception:
            return
        self.flush_para()
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        stream = io.BytesIO(blob)
        try:
            run.add_picture(stream, width=Inches(5.4))
        except Exception:
            return

    def flush_table(self) -> None:
        rows = self.table_rows or []
        self.table_rows = None
        if not rows:
            return
        cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                text, header = row[c_idx] if c_idx < len(row) else ("", False)
                cell = table.cell(r_idx, c_idx)
                cell.text = text
                if header or r_idx == 0 and any(item[1] for item in row):
                    set_cell_shading(cell, "1E3A8A")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def convert_file(src: Path, dst: Path) -> None:
    html = src.read_text(encoding="utf-8-sig")
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(2.0)
    parser = HtmlDocParser(document)
    parser.feed(html)
    parser.close()
    document.save(dst)


def main() -> None:
    docs = [
        path
        for path in sorted(PDFS.glob("*.doc"))
        if path.stat().st_ctime and datetime.fromtimestamp(path.stat().st_ctime).date() == TODAY
    ]
    if not docs:
        print("Nenhum .doc criado hoje em pdfs/")
        return

    print(f"Convertendo {len(docs)} arquivo(s) .doc de hoje…")
    for src in docs:
        dst = src.with_suffix(".docx")
        convert_file(src, dst)
        print(f"DOCX: {dst.name} ({dst.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
