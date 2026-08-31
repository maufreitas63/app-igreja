#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o Manual de Operação didático (Faça Você Mesmo) em .docx."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0F, 0x4C, 0x81)
MUTED = RGBColor(0x4A, 0x55, 0x68)
RED = RGBColor(0x8B, 0x1E, 0x1E)
GREEN = RGBColor(0x1B, 0x5E, 0x20)

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "MANUAL_OPERACAO_CONECTA_FACA_VOCE_MESMO.docx"


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run, *, size=11, bold=False, color=None, italic=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, color=None, italic=False, space_after=8, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_steps(doc, steps: list[str]) -> None:
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        n = p.add_run(f"{i}. ")
        set_run_font(n, size=11, bold=True, color=ACCENT)
        t = p.add_run(step)
        set_run_font(t, size=11)


def add_heading_custom(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else ACCENT
        run.font.name = "Calibri"


def add_label_block(doc, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(label)
    set_run_font(r, size=11, bold=True, color=NAVY)
    add_para(doc, body, size=11, space_after=6)


def add_feature(doc, *, code: str, name: str, what: str, route: str, who: str, access: list[str], manages: str, acl: str, fill: list[str], impact: str) -> None:
    add_heading_custom(doc, f"{code}  {name}", 2)
    add_label_block(doc, "Referência e nome da funcionalidade", f"{what}\nRota técnica: {route}")
    add_label_block(doc, "Quem acessa / visualiza", who)
    add_para(doc, "Como acessar (passo a passo)", size=11, bold=True, color=NAVY, space_after=4, space_before=6)
    add_steps(doc, access)
    add_label_block(doc, "Quem gerencia / cria / altera", manages)
    add_label_block(doc, "Permissão necessária (ACL)", acl)
    add_para(doc, "Como e quando preencher (guia prático)", size=11, bold=True, color=NAVY, space_after=4, space_before=6)
    add_steps(doc, fill)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("O que implica se não for preenchido a contento")
    set_run_font(r, size=11, bold=True, color=RED)
    add_para(doc, impact, size=11, space_after=14)


def add_page_numbers(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Conecta+  ·  Manual Faça Você Mesmo  ·  p. ")
    set_run_font(run, size=9, color=MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)
    set_run_font(r2, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    add_page_numbers(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    for i in range(1, 4):
        hs = styles[f"Heading {i}"]
        hs.font.color.rgb = NAVY
        hs.font.name = "Calibri"
        hs.font.bold = True


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "CONECTA+", size=14, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Manual de Operação e Guia Passo a Passo",
        size=26,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_para(
        doc,
        "Faça Você Mesmo — da tela inicial à engrenagem",
        size=16,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )
    add_para(
        doc,
        "Documento gerado a partir do código-fonte publicado (menus, rotas e ACL). Inclui Cantinho da Leitura, Livros doados (ISBN/Bipar/CBL e empréstimos), Declaração de Privacidade no Sobre o Conecta+, Modo Ghost e Aliança Conecta Reino. Papéis Líder, Líder Geral, Administrador de Eventos e Orquestrador de Evento foram absorvidos pela Secretaria e não existem mais.",
        size=12,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        f"Data de geração: {date.today().strftime('%d/%m/%Y')}",
        size=12,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_para(
        doc,
        "Ordem deste manual = ordem da tela: (1) Início, (2) Menu do membro, (3) Engrenagem / Configurações.",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    doc.add_page_break()


def add_howto(doc: Document) -> None:
    add_heading_custom(doc, "Como usar este manual", 1)
    add_para(
        doc,
        "Este é um guia de operação, não um folheto institucional. Siga a ordem das telas como se estivesse com o aplicativo aberto. Cada funcionalidade responde às mesmas sete perguntas, sempre nesta sequência.",
    )
    bullets = [
        "Referência e nome — o que é e qual a rota técnica.",
        "Quem acessa — quais papéis costumam ver o item (o Super Administrador vê tudo).",
        "Como acessar — o caminho exato na interface.",
        "Quem gerencia — quem cria, altera ou apaga.",
        "Permissão (ACL) — a chave conferida no banco (access_grants / sessionHasAccess).",
        "Como preencher — o passo a passo do dia a dia.",
        "Se não preencher — o impacto operacional.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r = p.add_run(b)
        set_run_font(r, size=11)

    add_heading_custom(doc, "Três gestos que você vai repetir o tempo todo", 2)
    add_steps(
        doc,
        [
            "Menu do membro: no canto superior esquerdo, toque no ícone de três traços (hambúrguer). Abre o painel «Menu».",
            "Engrenagem: ainda no Menu, no canto superior direito, toque na engrenagem. Só aparece se o seu papel tiver pelo menos um item de gestão. Título da tela: «Configurações».",
            "Fechar: na base da maioria das telas há a barra «Fechar». Use-a para voltar; não use o botão voltar do navegador como hábito.",
        ],
    )
    add_para(
        doc,
        "Itens que você não tem permissão simplesmente não aparecem. Se um colega vê a engrenagem e você não, a diferença é o papel (ACL), não um defeito da tela.",
        italic=True,
        color=MUTED,
    )

    add_heading_custom(doc, "Papéis do aplicativo (quem é quem)", 2)
    roles = [
        ("Visitantes", "visitantes", "Quem ainda não tem perfil completo. Acesso mínimo."),
        ("Congregado", "congregado", "Cadastrado, com acesso básico; sem gerir família nem finanças globais."),
        ("Membro", "member", "Acesso padrão da vida na igreja (Início, perfil, ofertas, célula, escalas…)."),
        ("Responsável familiar", "family_acceptor", "Complemento: gerencia integrantes da família."),
        ("Secretaria", "secretaria", "Operação da igreja: eventos, orquestração, escalas (todos os tipos), salas, totem, células, recepção, avisos, mídia, murais, campanhas, acervo/empréstimos de livros e temas da Trilha. Sem Cuidados Pastorais nem tesouraria global."),
        ("Tesoureiro", "tesoureiro", "Lançamentos, RD, orçamento, modelo preditivo; também opera campanhas."),
        ("Equipe Pastoral", "pastoral", "Membro + Cuidados Pastorais, mudança de papéis e temas da Trilha."),
        ("Gestor em Controle de Acesso", "gestor_controle_acesso", "Matriz de permissões. Nunca vê nem edita o Super Administrador nem PIN."),
        ("Super administrador", "super_admin", "Acesso pleno e global (curinga). Sem alteração neste perfil."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, t in enumerate(("Nome na tela", "Código ACL", "Para que serve")):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1E3A5F")
    for name, code, purpose in roles:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = code
        row[2].text = purpose
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()
    add_para(
        doc,
        "Regras que atravessam o manual: (1) cada igreja (tenant) só enxerga os próprios dados; (2) alguns menus exigem vínculo ativo (membership_out vazio); (3) no Modo Ghost a identidade usada é a da pessoa auditada, não a do operador; (4) o Gestor não lista Super Administrador; (5) não existem mais os papéis Líder, Líder Geral, Administrador de Eventos nem Orquestrador — a operação correspondente é da Secretaria; (6) desligado no acervo de livros é só membership_out, nunca o flag is_active do cadastro.",
        size=11,
        space_after=12,
    )
    doc.add_page_break()


def add_part0(doc: Document) -> None:
    add_heading_custom(doc, "0. Porta de entrada — antes da tela inicial", 1)
    add_para(
        doc,
        "Você só chega no Início depois de autenticar. Estes três passos não estão no menu lateral; são a porta.",
    )

    add_feature(
        doc,
        code="0.1",
        name="Login",
        what="Tela de entrada do Conecta+. Identifica a pessoa pelo telefone e pelo PIN de acesso (e biometria, se habilitada no dispositivo).",
        route="/  (app/index.tsx)",
        who="Qualquer pessoa com telefone cadastrado na instância. Sem login, o restante do aplicativo não abre.",
        access=[
            "Abra o endereço da igreja (PWA) ou o aplicativo.",
            "Informe o telefone com DDD (formato Brasil).",
            "Informe o PIN de acesso. Se esqueceu, use o envio de PIN por e-mail (quando o cadastro tiver e-mail).",
            "Se a igreja usar código de instância, preencha quando solicitado.",
            "Conclua. O sistema decide o próximo passo: cadastro incompleto, LGPD pendente ou Início.",
        ],
        manages="A pessoa altera o próprio PIN em Perfil → Dados Cadastrais. A Secretaria/Gestor não vê PIN. O Super Administrador opera a instância, não o PIN alheio pela matriz.",
        acl="Tela `/`. A sessão grava o telefone e resolve o perfil efetivo. Sem sessão válida, as rotas de membro redirecionam para o login.",
        fill=[
            "Use o mesmo telefone do cadastro (WhatsApp da família).",
            "PIN: o combinado na igreja; não compartilhe.",
            "Se a tela pedir biometria, autorize no dispositivo.",
        ],
        impact="Telefone ou PIN errados: você não entra. Sem e-mail no cadastro, a recuperação de PIN falha. Visitante só com telefone+PIN incompleto é mandado para o cadastro inicial.",
    )

    add_feature(
        doc,
        code="0.2",
        name="Cadastro inicial",
        what="Completa o perfil recém-criado: nome, nascimento, endereço (CEP) e, se a igreja exigir, aceite LGPD e selfie.",
        route="/register",
        who="Quem acabou de receber telefone+PIN e ainda não tem nome completo (fluxo «Visitante»).",
        access=[
            "Após o login, se o perfil estiver incompleto, o aplicativo abre esta tela sozinho.",
            "Não há atalho no Menu para «fazer o cadastro de novo».",
        ],
        manages="A própria pessoa. Depois, Secretaria pode ajustar dados em Cadastro de Usuário / Recepção Familiar.",
        acl="Tela `/register`. O aceite LGPD, quando o módulo está ativo na instância, é obrigatório para concluir.",
        fill=[
            "Preencha Nome completo (como no documento).",
            "Data de nascimento (usada em Aniversariantes e relatórios).",
            "CEP com 8 dígitos — o endereço é buscado automaticamente; confira número e complemento.",
            "Se aparecerem os termos LGPD: role até o fim, marque o aceite.",
            "Siga para a câmera, tire a selfie e confirme.",
        ],
        impact="Cadastro pela metade: a pessoa não chega no Início. Sem CEP, o mapa e a célula próxima não funcionam. Sem nascimento, some da lista de aniversariantes. Recusar LGPD (quando exigido) trava a conta.",
    )

    add_feature(
        doc,
        code="0.3",
        name="Aceite LGPD (perfil já existente)",
        what="Quando a igreja ativa o módulo LGPD, quem ainda não aceitou os termos passa por esta tela antes da home.",
        route="/lgpd",
        who="Membro/congregado com cadastro, enquanto `lgpd_accepted` não for verdadeiro e o parâmetro da instância estiver ligado.",
        access=["O redirecionamento é automático após o login. Não use o Menu."],
        manages="Governança da instância (parâmetro LGPD). O texto dos termos é da igreja.",
        acl="Tela `/lgpd` (`ACCESS_SCREEN.lgpd`).",
        fill=["Leia os termos, aceite e, se pedido, confirme a selfie."],
        impact="Sem aceite, a home não abre. A operação pastoral e de mídia fica juridicamente descoberta para aquela pessoa.",
    )

    add_feature(
        doc,
        code="0.4",
        name="Cadastro de família (formulário público)",
        what="Página pública para a família se inscrever antes (ou fora) do aplicativo. No PWA esta rota só encaminha para o formulário avulso.",
        route="/cadastro-familia  (e página standalone /cadastro-familia/)",
        who="Visitantes e famílias novas, sem login. O link costuma ser o QR da instância (Redes Sociais) ou o enviado pela recepção.",
        access=[
            "Abra o link público da igreja (não o Menu).",
            "Preencha o informante e os demais integrantes.",
            "Envie. A Secretaria processa o lote em Recepção Familiar.",
        ],
        manages="Secretaria (e Gestor / Super Admin) processa ou rejeita o lote. A família não «se aprova» sozinha.",
        acl="Fluxo público + painel `maintenance.card.profile_cadastro` para processar.",
        fill=[
            "Informante: nome, telefone, parentesco.",
            "Demais membros: nome e nascimento.",
            "Confira se todos da mesma casa estão no mesmo envio.",
        ],
        impact="Lote não enviado: a família não entra na fila. Lote enviado e não processado: não existe perfil no app, a régua de acolhimento não começa e o culto não tem inscrição familiar.",
    )
    doc.add_page_break()


def add_part1(doc: Document) -> None:
    add_heading_custom(doc, "1. Tela inicial (fora do menu lateral)", 1)
    add_para(
        doc,
        "Rota da home: /(tabs)  —  arquivo app/(tabs)/index.tsx. No Menu este destino se chama «Início». Não há mais carrossel de cards no Painel; o antigo /(tabs)/dashboard só redireciona. O que você vê abaixo está na própria tela, sem abrir o Menu.",
    )

    add_feature(
        doc,
        code="1.1",
        name="Topo da home (saudação, menu e logo)",
        what="Faixa superior da tela inicial: cumprimenta pelo primeiro nome, abre o Menu e mostra o logo da igreja.",
        route="Chrome compartilhado (MinimalScreenLayout + MinimalTopLeftChrome) sobre /(tabs)",
        who="Qualquer pessoa autenticada que chegou no Início.",
        access=[
            "Entre no aplicativo (login concluído).",
            "Você já está no Início. Não precisa abrir o Menu para ver eventos e avisos.",
        ],
        manages="O logo vem de Instâncias (Igrejas) — só Super Administrador. O nome vem dos Dados Cadastrais da própria pessoa.",
        acl="A home do Menu (`events_panel`) fica sempre habilitada após o login de membro.",
        fill=[
            "Confira se a saudação mostra o seu nome. Se aparecer «Visitante», complete o cadastro.",
            "Toque no hambúrguer (três traços, à esquerda) quando quiser o Menu.",
            "A engrenagem NÃO fica neste topo. Ela só aparece depois de abrir o Menu, se você tiver permissão de gestão.",
        ],
        impact="Nome errado na saudação = cadastro errado (corrija em Perfil). Sem logo, a identidade visual da igreja some do chrome, mas o app funciona.",
    )

    add_feature(
        doc,
        code="1.2",
        name="Próximos Eventos",
        what="Primeira página da caixa de entrada da home. Lista os próximos cultos/eventos publicados (nome, local, data e hora).",
        route="/(tabs) → EventsInboxHome (página Eventos)",
        who="Membro autenticado. Eventos só-membros respeitam o interruptor «Somente membros» da Programação de Eventos.",
        access=[
            "No Início, olhe o miolo da tela (não o Menu).",
            "Se estiver em Avisos, use a faixa «Avisos ↔ Eventos» até «Proximos Eventos».",
        ],
        manages="Secretaria, Gestor ou Super Admin em Engrenagem → Culto e Eventos → Programação de Eventos. Publicar o evento é o que o faz aparecer aqui.",
        acl="Os eventos vêm da tabela `events` filtrados pelo tenant. A home não exige card de carrossel. Manutenção: `maintenance.card.events`.",
        fill=[
            "Leia nome, local e horário. Se o local estiver vazio, a linha mostra «Sem local informado».",
            "Toque no evento para abrir a Agenda da Família (item 1.4).",
            "Se a lista estiver vazia, não há evento publicado à frente — avise a Secretaria, não «cadastre o culto» por aqui (membro não cria evento nesta tela).",
        ],
        impact="Sem eventos publicados: a home fica com «Nenhum evento disponível». Sem data/hora/local corretos na programação, a família marca presença no culto errado ou desiste de ir.",
    )

    add_feature(
        doc,
        code="1.3",
        name="Avisos da home",
        what="Segunda página da caixa de entrada. Cartões de comunicados publicados, oportunidades, generosidade, troca de escala, pastoral e campanhas.",
        route="/(tabs) → EventsInboxHome (página Avisos)",
        who="Membro autenticado. Cada tipo de aviso só aparece se existir conteúdo e se a pessoa tiver direito àquele módulo.",
        access=[
            "No Início, deslize ou toque a faixa até «Avisos».",
            "Toque no cartão para ir ao destino (mural, escala, pastoral, etc.).",
        ],
        manages="Secretaria: Engrenagem → Culto e Eventos → Manutenção de Avisos (comunicados). Outros avisos nascem dos respectivos módulos (vaga, troca, pedido pastoral).",
        acl="Comunicados: `maintenance.card.event_orchestration` e tabelas `event_avisos` / `event_control`. Ao abrir, vários tipos são marcados como lidos.",
        fill=[
            "Leia o aviso no próprio cartão.",
            "Se for um recado da liderança (Manutenção de Avisos), não há formulário para o membro — só leitura.",
            "Se for oportunidade/generosidade/troca, siga o botão do cartão e complete o fluxo da tela de destino.",
        ],
        impact="Sem avisos publicados e sem movimentos nos murais, a página mostra «Nenhum aviso publicado». A congregação não vê comunicados do culto nem convites operacionais.",
    )

    add_feature(
        doc,
        code="1.4",
        name="Agenda da Família",
        what="Tela sobreposta à home para a família confirmar quem vai ao evento (audiência). É o check-in familiar do culto, inclusive quórum e geofence quando o evento pede.",
        route="Sobreposta em /(tabs) (não é item do Menu). Abre ao tocar um evento.",
        who="A pessoa logada e os integrantes da família vinculados. Sem família, a lista de audiência fica vazia.",
        access=[
            "No Início, em Próximos Eventos, toque no culto desejado.",
            "O rodapé «Eu quero…» some enquanto a agenda está aberta.",
            "Para sair, use «Fechar» — você volta à lista de eventos.",
        ],
        manages="A família marca presença. A Secretaria programa o evento (capacidade, quórum, totem, geofence, salas). O totem no hall confirma quem chegou com QR.",
        acl="Inscrições na tabela `event_registrations`. Evento com «somente membros» restringe. Quórum/totem/geofence são flags do evento (`requer_quorum`, `totem_ativo`, `geofence_ativo`).",
        fill=[
            "Confira se o evento no topo é o culto certo (nome, local, horário).",
            "Na lista Audiência, marque o quadrado de cada pessoa da casa que vai. Desmarque quem não vai (exceto quando o totem já confirmou — aí o item fica cadeado).",
            "Se o culto tiver geofence ativo, respeite o banner de status: só confirma perto do templo, quando a igreja configurou assim.",
            "Não invente nomes nesta lista: quem não está na família precisa ser incluído em Perfil → Gerenciar Família (ou pela Secretaria na Recepção).",
        ],
        impact="Família não marcada: a igreja subestima a audiência, o «copo» de vagas mente, KIDS/TEENS não prepara sala, e assembleia com quórum fica incompleta. Marcar quem não foi gera lista falsa. Sem família vinculada, o QR da carteirinha também não existe.",
    )

    add_feature(
        doc,
        code="1.5",
        name="Eu quero… — Contribuir (Dízimos e Ofertas)",
        what="Atalho da home para gerar o Pix de dízimo/oferta. Não lança tesouraria: a pessoa paga no banco e a tesouraria concilia depois.",
        route="/ofertas  (aberta pelo rodapé da home)",
        who="Quem tem o card de ofertas (padrão do Membro e papéis que o herdam). Sem grant, o toque avisa «Você não tem permissão para abrir Dízimos e Ofertas.»",
        access=[
            "No Início, com a agenda fechada, olhe o rodapé «Eu quero…».",
            "Toque em «Contribuir».",
            "Toque em «Dízimos e Ofertas».",
        ],
        manages="Tesoureiro / Super Admin cadastram a chave Pix da igreja em Instâncias. O membro só informa o valor.",
        acl="`dashboard.card.offerings` (DRAWER_OFFERINGS_RESOURCE). Rota `/ofertas`.",
        fill=[
            "Informe o valor com centavos, como o texto da tela pede.",
            "Copie o Pix Copia e Cola.",
            "Cole no aplicativo do banco e conclua o pagamento.",
            "Não existe botão «já paguei» nesta tela — a conciliação é da tesouraria.",
        ],
        impact="Valor em branco: o app pede para informar o valor. Sem chave Pix na instância: «Chave PIX indisponível» — a oferta não sai e a tesouraria não recebe pelo app.",
    )

    add_feature(
        doc,
        code="1.6",
        name="Eu quero… — Campanhas e Projetos",
        what="Atalho da home para contribuir em campanha ativa (Pix já identificado com o projeto).",
        route="/ofertas  com campaignId ou campaignContribute=1",
        who="Mesmo grant de ofertas. Só funciona se existir campanha ativa.",
        access=[
            "Início → «Eu quero…» → «Contribuir» → «Campanhas e Projetos».",
            "Se houver uma campanha só, abre direto. Se houver várias, escolha o projeto.",
        ],
        manages="Secretaria, Tesoureiro, Pastoral ou Super Admin em Engrenagem → Finanças → Gestão de Campanhas (criar, meta, datas, status).",
        acl="View: `dashboard.card.offerings`. Campanhas no painel: `maintenance.finance.campaigns`.",
        fill=[
            "Escolha a campanha (se a tela pedir).",
            "Informe o valor e copie o Pix identificado.",
            "Pague no banco.",
        ],
        impact="Nenhuma campanha ativa: toast «Nenhuma campanha ativa no momento.» Sem meta/datas na gestão, o membro não consegue contribuir pelo app e a campanha não anda.",
    )

    add_feature(
        doc,
        code="1.7",
        name="Eu quero… — Pedido de oração",
        what="Formulário Coração Aberto: o membro envia um pedido à equipe pastoral (motivo, para quem, destino Intercessão ou Sigilo).",
        route="/pastoral",
        who="Membro e papéis que herdam o card pastoral. A fila de atendimento (Cuidados Pastorais) é só da Equipe Pastoral.",
        access=[
            "Início → rodapé «Eu quero…» → «Fazer um pedido de Oração».",
            "Também existe a mesma rota a partir de atalhos de aviso, quando houver.",
        ],
        manages="Quem pede: o membro. Quem tria, agenda e responde: Equipe Pastoral em Cuidados Pastorais. Motivos/submotivos são cadastro pastoral.",
        acl="Tela `/pastoral`. Card `dashboard.card.pastoral`. Painel interno: `maintenance.card.pastoral_care` (Secretaria não tem).",
        fill=[
            "Escolha o Motivo e a Situação (submotivo).",
            "Em «Este pedido é para», escolha: eu, família ou terceiros. Se não for você, preencha nome/parentesco.",
            "Escolha Encaminhar para (Intercessão ou Sigilo Pastoral).",
            "Escreva o pedido no campo «Seu pedido».",
            "Envie. Depois acompanhe em Perfil / histórico pastoral, se a tela oferecer.",
        ],
        impact="Campo obrigatório vazio: o app recusa com «Para enviar o pedido, preencha o campo …». Pedido não enviado não entra na fila. Destino errado (Sigilo vs Intercessão) expõe ou esconde o conteúdo da equipe errada.",
    )
    doc.add_page_break()


def add_part2(doc: Document) -> None:
    add_heading_custom(doc, "2. Menu principal (interação do membro)", 1)
    add_para(
        doc,
        "Ordem idêntica a APP_DRAWER_MENU_ITEMS em lib/appDrawerMenu.ts. Abra o hambúrguer. Só os itens liberados para o seu papel aparecem. No rodapé do Menu: Encerrar sessão / Sair do aplicativo.",
    )
    add_para(
        doc,
        "Caminho comum a todos os itens deste capítulo: Início → ícone de três traços (Menu) → toque no nome do item.",
        italic=True,
        color=MUTED,
    )

    add_feature(
        doc,
        code="2.1",
        name="Início",
        what="Volta para a tela inicial descrita no capítulo 1.",
        route="/(tabs)  — moduleKey events_panel",
        who="Todo usuário autenticado no fluxo de membro. Item sempre habilitado no drawer.",
        access=["Menu → Início."],
        manages="Não há cadastro neste item; ele só navega.",
        acl="Sem gate de card. Sempre enabled.",
        fill=["Use quando estiver perdido em outra tela e quiser eventos, avisos e «Eu quero…» de novo."],
        impact="Não se «preenche». Se você nunca volta ao Início, perde avisos novos e a agenda da semana.",
    )

    add_feature(
        doc,
        code="2.2",
        name="Perfil",
        what="Hub «Perfil & Identidade»: carteirinha com QR de check-in, dados cadastrais, família, trilha, reembolsos e Cantinho da Leitura (conforme grants).",
        route="/perfil  — moduleKey menu_perfil  — card dashboard.card.grouped_manage",
        who="Membro e papéis com o card grouped_manage. Opções internas aparecem uma a uma conforme a ACL de cada tela filha.",
        access=["Menu → Perfil."],
        manages="A própria pessoa edita o que a coluna permitir. Secretaria ajusta cadastro operacional. Pastoral não substitui o preenchimento do membro.",
        acl="Card `dashboard.card.grouped_manage`. Filhas: `/manage-profile`, `/manage-members`, `/trilha-discipulado`, `/expense-report`.",
        fill=[
            "Abra Perfil e escolha a ação (as seções 2.2.1 a 2.2.6 abaixo).",
            "Se aparecer «Nenhuma opção de perfil disponível», falta grant — fale com a liderança, não tente URL direta.",
        ],
        impact="Perfil vazio: carteirinha sem QR, mapa sem pin, célula sem CEP, aniversário sumido, RD impossível. A operação inteira da família pisa neste hub.",
    )

    add_feature(
        doc,
        code="2.2.1",
        name="Carteirinha Digital (QR de check-in)",
        what="Identidade da família no culto. O QR é o código da família para o Totem no hall — não é o QR das redes sociais.",
        route="Painel dentro de /perfil",
        who="Quem abre o hub Perfil e tem família vinculada.",
        access=["Menu → Perfil → Carteirinha Digital."],
        manages="A vinculação da família sai da Recepção Familiar / Gerenciar Família. O totem lê o QR; a Secretaria configura o totem e o evento (Totem ativo).",
        acl="Mesmo hub grouped_manage. Sem family_id, a carteirinha avisa para vincular código em Dados Cadastrais.",
        fill=[
            "Abra a carteirinha no domingo, com brilho da tela alto.",
            "Apresente o QR no leitor do hall (Totem de check-in).",
            "Não tire print para terceiros — o QR identifica a sua casa.",
        ],
        impact="Sem família: «Vincule um código de família…» e o totem não confirma. Família errada: check-in no núcleo de outra casa. Evento sem Totem ativo: o QR não registra presença.",
    )

    add_feature(
        doc,
        code="2.2.2",
        name="Dados Cadastrais",
        what="Ficha da pessoa: pessoais, contato, endereço, veículos e PIN.",
        route="/manage-profile  (também embutido no hub Perfil)",
        who="A própria pessoa (view/update das colunas liberadas). Congregado e Membro. Secretaria pode corrigir no painel Cadastro de Usuário.",
        access=["Menu → Perfil → Dados Cadastrais."],
        manages="Membro (autoatendimento) + Secretaria (cadastro operacional). Gestor não vê PIN.",
        acl="Tela `/manage-profile`. Edição campo a campo (`column` em access_resources, PROFILE_MANAGE_COLUMN_FIELDS).",
        fill=[
            "Dados pessoais: nome completo, CPF, nascimento.",
            "Contato: e-mail e telefone (e-mail é a recuperação de PIN).",
            "Endereço: CEP com 8 dígitos, aguarde o preenchimento, complete número e complemento. Alertas alimentares, se couber.",
            "Veículos: placa se você serve em estacionamento.",
            "PIN: altere só em local privado; confirme o PIN atual quando pedido.",
            "Salve cada campo quando a tela oferecer confirmação.",
        ],
        impact="Onboarding incompleto gera alerta na própria tela. Sem CEP: mapa e célula próxima falham. Sem e-mail: esqueceu o PIN e trava. Placa vazia: escala de estacionamento não identifica o carro. CPF/endereço errados quebram RD, mídia e correspondência.",
    )

    add_feature(
        doc,
        code="2.2.3",
        name="Gerenciar Família",
        what="Lista quem mora na mesma casa no aplicativo: incluir, tirar, foto, herança de endereço.",
        route="/manage-members",
        who="Membro e Responsável familiar (`family_acceptor`). Congregado em geral só vê, sem gerir.",
        access=["Menu → Perfil → Gerenciar Família."],
        manages="O responsável da casa. Secretaria processa famílias novas na Recepção Familiar; não deixe visitante «se promover» a membro por aqui.",
        acl="Tela `/manage-members`. Papel `family_acceptor` soma update. Isolamento por família/tenant.",
        fill=[
            "Confira se todos da casa estão na lista.",
            "Adicione integrante com nome e nascimento; herde endereço quando for o mesmo lar.",
            "Remova só quem realmente saiu da casa — a tela pede confirmação.",
            "Atualize foto se a igreja usar reconhecimento no totem.",
        ],
        impact="Filho de fora da lista: some da Agenda da Família e do quórum. Endereço divergente: o mapa cria pin errado. Remoção indevida: a pessoa perde o QR da casa.",
    )

    add_feature(
        doc,
        code="2.2.4",
        name="Trilha de Discipulado",
        what="Percurso de lições da igreja. A Lição 5.1 (Perfil Ministerial) alimenta o Mural de Oportunidades.",
        route="/trilha-discipulado",
        who="Congregado, Membro, lideranças e Secretaria (herança member). Conteúdo das lições é cadastrado na engrenagem.",
        access=["Menu → Perfil → Trilha de Discipulado."],
        manages="Secretaria e Equipe Pastoral: Temas da Trilha. Super Admin: Resetar Trilha. Certificados: Reconhecimentos.",
        acl="Tela `/trilha-discipulado`. Painéis: `maintenance.card.discipleship_themes`, `_alerts`, `_reset` (reset só Super Admin).",
        fill=[
            "Abra a lição da vez, leia texto, assista ao vídeo se houver, registre a reflexão quando a lição pedir.",
            "Na Lição 5.1, preencha o Perfil Ministerial com honestidade (dons, disponibilidade).",
            "Não pule o envio: progresso só conta o que a trilha marcar como concluído.",
        ],
        impact="Trilha vazia de temas: o membro abre uma trilha sem conteúdo. Sem 5.1: o Mural de Oportunidades bloqueia com «abra a Trilha». Reset indevido (só Super Admin) apaga progresso real.",
    )

    add_feature(
        doc,
        code="2.2.5",
        name="Reembolsos (Relatório de Despesas)",
        what="O membro envia comprovantes de despesa feita em nome da igreja (RD) para a tesouraria.",
        route="/expense-report",
        who="Membro com grant de RD. Tesoureiro recebe, vincula e concilia em Informações Financeiras.",
        access=["Menu → Perfil → Reembolsos."],
        manages="Quem gastou preenche. Tesoureiro emite/vincula/concili. Super Admin vê tudo.",
        acl="Tela `/expense-report`. Painel tesouraria: `maintenance.card.financials`.",
        fill=[
            "Crie um novo RD.",
            "Informe cada despesa (valor, descrição) e anexe comprovante nítido.",
            "Envie. Acompanhe o status; não «lance» você mesmo no financeiro da igreja.",
        ],
        impact="Sem descrição/comprovante: a tesouraria não concilia. RD não enviado = a igreja não devolve e o livro caixa fica incompleto. Acesso negado: toast na tela.",
    )

    add_feature(
        doc,
        code="2.2.6",
        name="Cantinho da Leitura",
        what="Reserva de livros do acervo da igreja: lista alfabética (um título por linha), ficha (capa, autor, editora, ano, ISBN) e prazo de 30 dias. A Secretaria confirma a retirada.",
        route="Painel MeusLivrosRetiradosPanel dentro de /perfil",
        who="Quem abre o hub Perfil (card grouped_manage). O acervo é o mesmo de Livros doados.",
        access=["Menu → Perfil → Cantinho da Leitura."],
        manages="O membro reserva e pode cancelar a reserva. Secretaria confirma retirada, renova (+10 dias), registra devolução e cadastra o ISBN em Livros doados.",
        acl="Hub `dashboard.card.grouped_manage`. Operação do acervo: tela `/livros-doados`. RPCs `list_livros_disponiveis_reserva`, `reservar_livro_acervo`, `cancelar_reserva_livro`. Isolamento por tenant.",
        fill=[
            "Abra a lista (ordem A–Z). Toque no título para ver capa, autor, editora, ano e ISBN.",
            "Escolha a data de retirada. O retorno inicial é 30 dias.",
            "Toque em Reservar — o botão só habilita com um livro selecionado.",
            "Acompanhe «Meus empréstimos». Cancele a reserva se desistir antes da Secretaria confirmar.",
            "Não cadastre ISBN aqui: doação e bipar são da Secretaria (item 3.B.12).",
        ],
        impact="Sem acervo cadastrado: lista vazia. Reservar sem confirmação da Secretaria: o livro não sai. Duas famílias na mesma reserva: o índice único do livro ativo impede — o título some da lista disponível.",
    )

    add_feature(
        doc,
        code="2.3",
        name="Financeiro",
        what="Leitura dos boletins da igreja (mês, comparativo, 12 meses, orçamento, saldo). Não é tela de lançar despesa.",
        route="/financial  — moduleKey gestao_financeira",
        who="Quem tem `/financial` / card `dashboard.card.financial` (Membro em view típico; Tesoureiro opera o painel de manutenção). Secretaria NÃO herda este card.",
        access=["Menu → Financeiro."],
        manages="Tesoureiro lança e fecha o mês em Informações Financeiras. O membro só consulta.",
        acl="`ACCESS_SCREEN.financial` = `/financial` + `dashboard.card.financial`.",
        fill=[
            "Escolha o mês de referência no alto.",
            "Abra as seções em acordeão: Resumo, Resultado, Comparativo, 12 meses, Histórico, Planejado × Realizado, Saldo.",
            "Se «Planejado × Realizado» estiver bloqueado, ainda não há orçamento naquele mês — cobrado à tesouraria, não preenchido aqui.",
            "Se a carga falhar, use «Atualizar».",
        ],
        impact="Tesouraria sem lançamentos: o membro vê vazio ou erro. Sem orçamento: não há planejado×realizado. Transparência da assembleia cai.",
    )

    add_feature(
        doc,
        code="2.4",
        name="Minha Célula",
        what="Pequeno grupo da pessoa: encontrar grupo próximo, pedir para entrar/sair, ver roteiro da semana e avisar ausência.",
        route="/pequeno-grupo  — card dashboard.card.small_group",
        who="Membro com vínculo ativo (`membership_out` vazio) e grant do card. Visitante sem membership não vê o item.",
        access=["Menu → Minha Célula."],
        manages="Secretaria e Pastoral cadastram grupos em Gestão de Pequenos Grupos. O membro pede vaga; o anfitrião ou o líder da célula recebe o fluxo.",
        acl="`dashboard.card.small_group`. Admin: `maintenance.card.small_groups_management` (update = pode administrar). Exige membership ativo.",
        fill=[
            "Se ainda não tem grupo: cadastre CEP no Perfil. Veja anfitriões próximos e toque «Quero participar».",
            "Se já tem grupo: confira horário e endereço. Abra a lista de participantes e o Roteiro da Semana.",
            "Faltar: «Avisar Ausência» (WhatsApp ao líder). Sem telefone do líder, a mensagem falha — avise a Secretaria para completar o cadastro do líder.",
            "Sair: «Quero sair do Grupo» (anfitrião/líder não usam este botão da mesma forma).",
        ],
        impact="Sem CEP: não há distância nem convite geográfico. Sem grupos cadastrados: lista vazia e a régua de acolhimento (dia 4) não tem para onde convidar. Sem roteiro: o encontro da semana fica no improviso.",
    )

    add_feature(
        doc,
        code="2.5",
        name="Escalas",
        what="O servo vê em que culto está escalado, pede troca (se permitido) e, em alguns tipos, informa placa ou fala com a equipe.",
        route="/escalas  — card dashboard.card.vigilance_scales",
        who="Membro com vínculo ativo e card de escalas. Tipos de escala podem ser filtrados por ACL (`scale_type.*`). Troca: `scales.allow_swap`.",
        access=["Menu → Escalas. Use as abas Escalas e Pedidos de troca."],
        manages="Secretaria programa tipos, servos e datas (acesso a todas as escalas ativas). O membro só solicita troca.",
        acl="`dashboard.card.vigilance_scales`. Painéis: `maintenance.card.scale_types`, `scale_volunteers`, `scales`. Estacionamento: `dashboard.card.parking_vehicle_v2`.",
        fill=[
            "Selecione o tipo de escala (Louvor, Recepção, etc.). Se o tipo não abrir, você não lidera/serve naquele tipo.",
            "Veja a data em que está escalado.",
            "Se precisar faltar e a igreja permitir: «Solicitar troca desta data» e acompanhe a aba Pedidos de troca.",
            "Estacionamento: mantenha a placa em Dia Cadastrais / no painel da escala.",
            "Intercessão: use o WhatsApp da lista quando a tela oferecer.",
        ],
        impact="Programação vazia: o membro acha que não serve e o culto fica descoberto. Troca sem grant: o botão não resolve. Placa incompleta: o estacionamento não identifica o carro. Sem tipos cadastrados, ninguém consegue escalar.",
    )

    add_feature(
        doc,
        code="2.6",
        name="Mural de Oportunidades",
        what="Vagas de serviço com matching do Perfil Ministerial (Lição 5.1). O membro declara interesse; o líder recebe.",
        route="/mural-oportunidades  — dashboard.card.opportunities",
        who="Membro com vínculo ativo e card de oportunidades.",
        access=["Menu → Mural de Oportunidades."],
        manages="Liderança cria vagas em Mural de Voluntários (engrenagem). Matching usa a Trilha.",
        acl="`dashboard.card.opportunities`. Admin: `maintenance.volunteer.mural`. Membership ativo obrigatório.",
        fill=[
            "Se a tela mandar abrir a Trilha, conclua a Lição 5.1 primeiro.",
            "Olhe «Vagas para você» (selo Match %) e «Outras vagas».",
            "Toque «Tenho interesse». O app registra e tenta WhatsApp ao líder.",
            "Não publique vaga por aqui — isso é tela de interesse, não de cadastro de vaga.",
        ],
        impact="Sem 5.1: mural bloqueado. Sem vagas no painel admin: ninguém serve. Interesse sem WhatsApp do líder: o registro existe, o contato humano falha.",
    )

    add_feature(
        doc,
        code="2.7",
        name="Mural de Generosidade",
        what="Doações e pedidos de empréstimo entre irmãos, com moderação. O telefone do anunciante não vai para o feed público.",
        route="/mural-generosidade",
        who="Membro com vínculo ativo e grant `/mural-generosidade` (e/ou card).",
        access=["Menu → Mural de Generosidade."],
        manages="O membro publica (vai para moderação). Secretaria, Pastoral ou Super Admin aprovam em Moderação do Mural.",
        acl="Tela `/mural-generosidade`. Moderação: `maintenance.card.generosity_moderation`. Membership ativo.",
        fill=[
            "Nova publicação: escolha Doar ou Pedir empréstimo.",
            "Categoria (Móveis, Saúde, Vestuário, Livros, Outros), Título e Descrição. Foto ajuda, mas não substitui o texto.",
            "«Enviar para moderação» — não aparece no mural até aprovarem.",
            "Nas abas, use «Tenho interesse» / «Posso ajudar» (o nome vai à liderança, não o telefone no mural).",
            "Quando resolver, «Marcar resolvido». Acompanhe «Meus anúncios».",
        ],
        impact="Título/descrição vazios: não publica. Sem moderação: fila eterna, mural público parado. Anúncio resolvido não marcado: outra família perde tempo. Nunca coloque telefone no texto — a regra do produto é não expor WhatsApp no feed.",
    )

    add_feature(
        doc,
        code="2.8",
        name="Sugestões",
        what="Canal para registrar melhorias e problemas do aplicativo/operação, com acompanhamento de status.",
        route="/suggestions-improvements",
        who="Membro ativo com card Administrativo, ou quem tem o painel de manutenção de sugestões. Liderança com `/maintenance-dashboard` também.",
        access=["Menu → Sugestões."],
        manages="Quem sugere: o membro. Quem trata: equipe com grant `maintenance.card.suggestions_improvements` (update para gerir).",
        acl="`isSuggestionsImprovementsAccessAllowed`: painel `suggestions_improvements` OU (`dashboard.card.administrativo` + membership ativo).",
        fill=[
            "«Registrar solicitação».",
            "Tipo de registro e Tema (opcional).",
            "Descrição detalhada — obrigatória. Escreva o que tentou fazer e o que esperava.",
            "Ligue notificações no app e/ou WhatsApp se quiser retorno.",
            "Anexe imagem se ajudar (print da tela).",
            "Envie e acompanhe a lista/detalhe.",
        ],
        impact="Sem descrição: «Informe a descrição detalhada…». Sugestão não enviada não entra na fila. Sem alguém com grant de gestão, as solicitações envelhecem sem dono.",
    )

    add_feature(
        doc,
        code="2.9",
        name="Redes Sociais",
        what="Atalhos para site, Instagram e YouTube da igreja, mais o QR da URL pública da instância (para divulgar o app, não para check-in).",
        route="/redes-sociais  — sempre visível no Menu",
        who="Qualquer pessoa no Menu do membro.",
        access=["Menu → Redes Sociais."],
        manages="Super Administrador em Instâncias (Igrejas): preenche as URLs. Sem URL, o botão nem aparece.",
        acl="Item sem gate de card no drawer (`menu_redes_sociais` sempre enabled).",
        fill=[
            "Toque em Site / Instagram / YouTube se existirem.",
            "Mostre o QR da instância para um visitante instalar/abrir o Conecta+.",
            "Não use este QR no totem do culto.",
        ],
        impact="URLs vazias: «Esta instância ainda não cadastrou site, Instagram ou YouTube.» A divulgação do app e das redes para. QR de instância ≠ QR da carteirinha.",
    )

    add_feature(
        doc,
        code="2.10",
        name="Sobre o Conecta+",
        what="Texto institucional, número de versão/revisão e a Declaração de Privacidade e Segurança de Dados (LGPD) em leitura completa.",
        route="/sobre-conecta  — sempre visível",
        who="Qualquer pessoa no Menu.",
        access=["Menu → Sobre o Conecta+."],
        manages="Equipe de produto/Super Admin (texto, versão e declaração no código). O membro só lê. O aceite jurídico de cadastro continua em /lgpd — esta tela não substitui o rito de aceite.",
        acl="Sempre enabled (`menu_sobre_conecta`).",
        fill=[
            "Leia o texto Sobre e anote Versão / Revisão se for abrir um chamado em Sugestões.",
            "Toque em «Declaração de Privacidade e Segurança de Dados (LGPD)».",
            "Role os sete pontos (isolamento por igreja, mural sem telefone, sem vitrine de dons, QR só com código familiar, mapa, sigilo pastoral, escudo do Super Administrador).",
            "Use Fechar no rodapé da declaração para voltar ao Sobre.",
        ],
        impact="Não há formulário. Sem esta tela, suporte e treinamento perdem a referência de versão e a congregação não encontra a declaração institucional no app.",
    )
    doc.add_page_break()


def add_part3(doc: Document) -> None:
    add_heading_custom(doc, "3. Engrenagem — Configurações (orquestração e gestão)", 1)
    add_para(
        doc,
        "Caminho único: Menu (hambúrguer) → ícone de engrenagem (só se canAccessSettings) → tela «Configurações». Os grupos são acordeões. A ordem visual abaixo é a de APP_DRAWER_SETTINGS_GROUPS + o rearranjo de AppDrawerSettings.tsx (Trilha agrupada, Assinaturas depois da Trilha, Instâncias fixas no rodapé).",
    )
    add_para(
        doc,
        "Quase todos os painéis exigem também a tela `/maintenance-dashboard` (view). Painéis internos abrem `/maintenance-dashboard?panel=CHAVE`. Ferramentas de operador (salas, totem, mídia, billing, igrejas) têm rota própria. Super Administrador vê tudo. Secretaria opera o culto e as pessoas, mas não vê Cuidados Pastorais, tesouraria global, Controle de Acesso, Ghost nem Assinaturas.",
        space_after=12,
    )

    add_heading_custom(doc, "3.A  Operação e Segurança", 1)

    add_feature(
        doc,
        code="3.A.1",
        name="Configuração de salas",
        what="Cadastro dos espaços físicos com nome afetivo e atribuição de membros (quem «mora» em cada sala no culto).",
        route="/configuracao-salas  — hint: Nomes afetivos e atribuição de membros",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Operação e Segurança → Configuração de salas."],
        manages="Secretaria (operação diária). Super Admin em última instância.",
        acl="Tela `/configuracao-salas`. Tabelas `church_room_settings` e `user_room_assignment`. Drawer: canManageRooms ou Super Admin.",
        fill=[
            "Aba Salas: crie a sala com o nome que a igreja usa no domingo (não o código interno).",
            "Salve. Informe vigência se a tela pedir data.",
            "Aba Membros: busque a pessoa e vincule à sala certa (KIDS, TEENS, etc.).",
            "Aba Distribuição: confira se ninguém ficou de fora do lote.",
            "Na Programação de Eventos, habilite as mesmas chaves de sala no culto.",
        ],
        impact="Sem salas: o check-in infantil não tem para onde ir e a tela «Sala(s) - Check In» opera no vazio. Membro sem atribuição: a criança não aparece para o servidor da sala. Nome genérico demais: a família não reconhece a porta.",
    )

    add_feature(
        doc,
        code="3.A.2",
        name="Totem de check-in",
        what="Leitor de QR no hall. Confirma que a família chegou ao culto do dia.",
        route="/totem-checkin  — hint: Leitor de QR no hall",
        who="Secretaria e quem tiver manutenção ou salas. O dispositivo do hall fica logado nesta tela o culto inteiro.",
        access=["Menu → engrenagem → Operação e Segurança → Totem de check-in. Use um tablet/celular fixo no hall."],
        manages="Secretaria configura. Em Programação de Eventos, ligue «Totem ativo» e publique o culto. Famílias geram o QR na Carteirinha.",
        acl="Tela `/totem-checkin`. Drawer: manutenção OU salas OU Super Admin. Evento precisa `totem_ativo`.",
        fill=[
            "No culto de hoje, confirme na Programação de Eventos: publicado + Totem ativo.",
            "Abra o Totem no aparelho do hall, autorize a câmera, deixe a tela ligada.",
            "Peça à família para abrir Perfil → Carteirinha e apontar o QR.",
            "Não use o Totem para «testar» QR de outra igreja/tenant.",
        ],
        impact="Evento sem Totem ativo: o QR não confirma. Câmera negada: fila no hall. Família sem QR: volta ao cadastro. Totem no perfil errado: presença lançada em outra casa.",
    )

    add_feature(
        doc,
        code="3.A.3",
        name="Autorização de imagem e voz",
        what="Termo LGPD de mídia (foto, vídeo, voz) com confirmação por e-mail e PDF.",
        route="/autorizacao-midia  — hint: Termos LGPD e confirmação por e-mail",
        who="Secretaria, Equipe Pastoral e Super Admin.",
        access=["Menu → engrenagem → Operação e Segurança → Autorização de imagem e voz."],
        manages="Secretaria ou pastoral aplica o termo. A pessoa titular confirma o e-mail.",
        acl="Tela `/autorizacao-midia`. Drawer: Super Admin ou grant `/autorizacao-midia`.",
        fill=[
            "Preencha nome, e-mail, CPF e telefone iguais ao cadastro.",
            "Leia o termo até o fim e aceite.",
            "Envie e peça à pessoa para confirmar o e-mail.",
            "Guarde/baixe o PDF quando a tela oferecer.",
        ],
        impact="Sem termo: a igreja não tem base para gravar o culto com aquela pessoa em close. E-mail errado: confirmação nunca chega.",
    )

    add_heading_custom(doc, "3.B  Gestão de Pessoas", 1)
    add_para(
        doc,
        "Lista de Membros, Mapa, Aniversariantes e Administrativo, no drawer, ainda exigem visão de manutenção (ou Super Admin), além do grant do card. Cuidados Pastorais é exclusivo da Equipe Pastoral.",
        italic=True,
        color=MUTED,
    )

    add_feature(
        doc,
        code="3.B.1",
        name="Lista de Membros",
        what="Diretório da comunidade (membros, famílias, visitantes) com busca e atalho de WhatsApp.",
        route="/membros  — hint: Diretório da comunidade  — card dashboard.card.members_list",
        who="Equipe Pastoral, Secretaria e Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Lista de Membros."],
        manages="Cadastros nascem da Recepção / Cadastro de Usuário / autoatendimento. Esta tela é consulta, não o formulário principal de alta.",
        acl="`dashboard.card.members_list`. Membership ativo + manutenção no drawer de settings.",
        fill=[
            "Busque pelo nome.",
            "Abra o núcleo familiar se precisar do contexto da casa.",
            "Use WhatsApp só para cuidado legítimo — não extraia lista para fora do app.",
        ],
        impact="Cadastros podres na origem: o diretório mente. Sem este grant, ninguém acha gente no domingo pelo diretório.",
    )

    add_feature(
        doc,
        code="3.B.2",
        name="Mapa de geolocalização",
        what="Pins das famílias (e células) no mapa da cidade.",
        route="/mapa-geolocalizacao  — hint: Pins das famílias no mapa",
        who="Quem tem a tela `/mapa-geolocalizacao` e membership ativo, e no menu da engrenagem também manutenção (ou Super Admin). Pastoral e Secretaria são o uso típico.",
        access=["Menu → engrenagem → Gestão de Pessoas → Mapa de geolocalização."],
        manages="Ninguém «desenha» o pin: ele nasce do CEP/endereço do cadastro. Secretaria cobra CEP na recepção.",
        acl="`ACCESS_SCREEN.mapGeolocation` = `/mapa-geolocalizacao`. Detalhe do pin tem tela extra.",
        fill=[
            "Espere o mapa carregar.",
            "Toque no pin para ver a família.",
            "Se faltar pin, vá ao Cadastro de Usuário e complete CEP (não «crie pin» solto).",
        ],
        impact="Sem CEP geocodificado: mapa vazio, célula próxima e visita pastoral geográfica falham.",
    )

    add_feature(
        doc,
        code="3.B.3",
        name="Aniversariantes",
        what="Quem faz aniversário, para cumprimento pastoral (WhatsApp).",
        route="/aniversariantes  — card dashboard.card.birthdays",
        who="Equipe Pastoral, Secretaria e Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Aniversariantes."],
        manages="A data vem dos Dados Cadastrais / recepção. Aqui só se consulta.",
        acl="`dashboard.card.birthdays`. Membership + manutenção no drawer.",
        fill=["Abra a lista da semana/mês. Cumpra pelo WhatsApp. Não altere nascimento nesta tela."],
        impact="Nascimento vazio no cadastro: a pessoa some da lista e ninguém parabeniza.",
    )

    add_feature(
        doc,
        code="3.B.4",
        name="Cuidados Pastorais",
        what="Fila de pedidos de oração (Intercessão e Sigilo), status e agenda de atendimentos.",
        route="/maintenance-dashboard?panel=pastoral_care  — hint: Fila e slots",
        who="Somente Equipe Pastoral e Super Admin. Secretaria não tem grant nem vê o item.",
        access=["Menu → engrenagem → Gestão de Pessoas → Cuidados Pastorais."],
        manages="Equipe Pastoral (triagem, status, slots). O membro só envia o pedido pelo «Eu quero…».",
        acl="`maintenance.card.pastoral_care` (+ `maintenance.pastoral.agenda`). Sigilo não aparece para voluntário de intercessão.",
        fill=[
            "Abra Pedidos. Separe Intercessão de Sigilo Pastoral.",
            "Atualize o status quando assumir o caso.",
            "Use Minha Agenda / slots para marcar conversa.",
            "Nunca encaminhe conteúdo de Sigilo para grupo aberto.",
        ],
        impact="Fila parada: o membro acha que ninguém orou. Agenda vazia: não há retorno. Secretaria tentando «ajudar» aqui seria vazamento — o ACL impede de propósito.",
    )

    add_feature(
        doc,
        code="3.B.5",
        name="Gestão de Pequenos Grupos",
        what="Administração de células: nome, horário, membros, visitantes, roteiro da semana, oração.",
        route="/maintenance-dashboard?panel=small_groups_management",
        who="Secretaria, Pastoral, Gestor, Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Gestão de Pequenos Grupos."],
        manages="Liderança da célula (anfitrião/líder do grupo) no dia a dia; Secretaria cadastra o grupo.",
        acl="`maintenance.card.small_groups_management` com update libera `can_admin_small_groups`.",
        fill=[
            "Novo grupo: nome que a igreja usa + horário HH:MM.",
            "Vincule líder, anfitrião e membros (telefones certos para o aviso de ausência).",
            "Toda semana: Roteiro (título, conteúdo, link).",
            "Registre visitantes da noite, se a tela oferecer.",
            "Salve. Confira se o endereço do anfitrião está no cadastro dele.",
        ],
        impact="Sem grupos: Minha Célula e o dia 4 da régua de acolhimento ficam sem destino. Sem roteiro: o encontro esvazia. Horário errado: a família vai na hora morta.",
    )

    add_feature(
        doc,
        code="3.B.6",
        name="Mural de Voluntários",
        what="Cadastro das vagas que alimentam o Mural de Oportunidades do membro (título, tipo de escala, status, interesses).",
        route="/maintenance-dashboard?panel=volunteer_mural",
        who="Secretaria, Pastoral, Gestor, Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Mural de Voluntários."],
        manages="Secretaria (ou pastoral) cadastra a vaga; o ministério dono trata os interesses.",
        acl="`maintenance.volunteer.mural`.",
        fill=[
            "Nova vaga: título claro («Violão no culto das 18h»).",
            "Associe o tipo de escala correspondente.",
            "Deixe status aberto quando quiser candidatos.",
            "Trate os interesses (WhatsApp / conversa) e feche a vaga quando preencher.",
        ],
        impact="Sem vagas: o mural do membro fica vazio mesmo com Lição 5.1 feita. Vaga sem tipo de escala: o matching e a programação não conversam.",
    )

    add_feature(
        doc,
        code="3.B.7",
        name="Moderação do Mural",
        what="Aprovar ou rejeitar doações e pedidos de empréstimo antes de irem ao mural público.",
        route="/maintenance-dashboard?panel=generosity_moderation  — hint: Doações e pedidos de empréstimo",
        who="Secretaria, Pastoral, Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Moderação do Mural."],
        manages="Secretaria ou pastoral. O membro só envia para moderação.",
        acl="`maintenance.card.generosity_moderation`.",
        fill=[
            "Fila Pendentes: abra o anúncio, leia título/descrição/foto.",
            "Aprove o que for seguro e cristão; rejeite o que for venda disfarçada ou dado demais.",
            "Acompanhe Publicados e Interesses — não deixe combinação sem dono.",
        ],
        impact="Fila sem moderador: nada novo no mural. Aprovar sem ler: risco a famílias. Rejeitar sem critério: desanima quem quer doar.",
    )

    add_feature(
        doc,
        code="3.B.8",
        name="Recepção Familiar",
        what="Fila dos formulários públicos de família. Processar cria perfis/membros; rejeitar descarta o lote. É a porta de entrada operacional da Secretaria.",
        route="/maintenance-dashboard?panel=family_reception",
        who="Secretaria, Gestor, Super Admin (grant `profile_cadastro`).",
        access=["Menu → engrenagem → Gestão de Pessoas → Recepção Familiar."],
        manages="Secretaria (moderação e aprovação). Não auto-promova visitante a Membro neste passo — o papel sobe depois, pela pastoral, se for o caso.",
        acl="O mesmo recurso de Cadastro de Usuário: `maintenance.card.profile_cadastro`. RPCs de recepção exigem este view.",
        fill=[
            "Atualize a fila. Abra o lote (informante + integrantes).",
            "Confira telefones e se a casa já existe (conflito de família).",
            "Processar se os dados estão bons; Rejeitar se for teste/duplicata, com critério.",
            "Anote o código familiar gerado. Oriente a família a fazer o primeiro login (telefone + PIN combinado).",
            "No mesmo dia, se couber, inicie a Régua de Acolhimento.",
        ],
        impact="Lote parado: a família não existe no app, não tem QR, não entra na agenda. Processar duplicado: duas casas para a mesma gente. Promover papel cedo demais foge da regra pastoral.",
    )

    add_feature(
        doc,
        code="3.B.9",
        name="Régua de Acolhimento",
        what="Tarefas da equipe de boas-vindas: WhatsApp no dia 1 e convite à célula no dia 4 (depois pode virar pendência pastoral).",
        route="/maintenance-dashboard?panel=visitor_followup  — hint: Tarefas da equipe de boas-vindas",
        who="Secretaria (welcome) e Super Admin. Pastoral vê o desdobramento se o caso virar cuidado.",
        access=["Menu → engrenagem → Gestão de Pessoas → Régua de Acolhimento."],
        manages="Equipe de recepção/Secretaria nas tarefas 1 e 4. Pastoral no transbordo.",
        acl="`maintenance.card.visitor_followup` (welcome também aceita `profile_cadastro` em RPCs).",
        fill=[
            "Atualize a lista do dia.",
            "Toque no WhatsApp, cumpra o roteiro (boas-vindas / convite à célula).",
            "Marque Concluído só depois de falar de verdade.",
            "Se não houver célula cadastrada, primeiro crie o grupo (3.B.5) — senão o dia 4 não tem destino.",
        ],
        impact="Tarefa não cumprida: o visitante esfria. Marcar concluído sem contato: a pastoral acha que já houve cuidado. Dia 8 sem fechamento vira peso pastoral.",
    )

    add_feature(
        doc,
        code="3.B.10",
        name="Cadastro de Usuário",
        what="Busca e correção pontual de pessoa já existente (nome, CEP, endereço). Mesmo grant da Recepção — por isso a Secretaria vê os dois itens.",
        route="/maintenance-dashboard?panel=profile_cadastro",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Cadastro de Usuário."],
        manages="Secretaria (higiene cadastral). Excluir usuário é irreversível no fluxo da tela — use com critério extremo.",
        acl="`maintenance.card.profile_cadastro`.",
        fill=[
            "Busque pelo nome ou telefone.",
            "Corrija nome, CEP, número, complemento.",
            "Salve. Confira o pin no mapa depois.",
            "Não use exclusão como «limpeza de fila» — rejeite na Recepção enquanto for lote pendente.",
        ],
        impact="CEP errado: mapa e célula quebram. Nome errado: lista, aniversário e Pix humano falham. Exclusão indevida: some histórico de presença e família.",
    )

    add_feature(
        doc,
        code="3.B.11",
        name="Administrativo",
        what="Atos constitutivos e documentos da igreja para consulta da comunidade autorizada.",
        route="/administrativo  — hint: Atos constitutivos  — card dashboard.card.administrativo",
        who="Quem tem o card administrativo; no menu da engrenagem, ainda precisa de manutenção ou Super Admin.",
        access=["Menu → engrenagem → Gestão de Pessoas → Administrativo."],
        manages="Super Admin / governança da instância sobe os documentos. O membro autorizado lê.",
        acl="`dashboard.card.administrativo`.",
        fill=["Abra o documento necessário (estatuto, ata). Não edite texto jurídico por improvisação — substitua o arquivo pela governança."],
        impact="Pasta vazia: assembleia sem base no app. Documento velho: decisão da igreja apoia-se em texto errado.",
    )

    add_feature(
        doc,
        code="3.B.12",
        name="Livros doados",
        what="Acervo da igreja (ISBN, Bipar, cadastro manual) e empréstimos da Secretaria: confirmar retirada, renovar +10 dias, devolver, WhatsApp do prazo e histórico.",
        route="/livros-doados  — hint: Acervo com busca ISBN e cadastro manual  — abas Acervo | Empréstimos | Histórico",
        who="Secretaria e Super Administrador com grant da tela `/livros-doados`. Gestor não opera o acervo por este item salvo grant explícito.",
        access=["Menu → engrenagem → Gestão de Pessoas → Livros doados."],
        manages="Secretaria cadastra doação e registra saída/devolução. O membro só reserva no Cantinho da Leitura (2.2.6).",
        acl="`ACCESS_SCREEN.livrosDoados` = `/livros-doados`. Tabelas `livros` e `emprestimos_livros`. RPCs em scripts/livros-doados.sql e scripts/emprestimos-livros.sql. Isolamento por tenant_id.",
        fill=[
            "Aba Acervo: abra «Registro de Doações». Digite o ISBN ou toque em Bipar (autorize a câmera). Confira título, autor, editora, ano e capa. Se a Google Books recusar (cota), o app tenta o cadastro brasileiro (CBL/BrasilAPI) e a Open Library. Complete o que faltar e salve. A seção fecha e o Acervo abre em ordem alfabética.",
            "Cadastro manual: use quando o ISBN não existir em nenhum catálogo.",
            "Aba Empréstimos: busque o membro por nome (mínimo 2 letras) ou telefone (mínimo 2 dígitos — senão a busca não lista a igreja inteira). Escolha o livro livre ou um título externo.",
            "Registrar empréstimo (30 dias) ou, se houver reserva, Confirmar retirada (use o diálogo da tela; no navegador não basta o alerta nativo).",
            "Nome em vermelho «(desligado)» só se membership_out estiver preenchido — não confundir com cadastro inativo (is_active).",
            "WhatsApp verde: mensagem de prazo e pergunta se devolve ou estende 10 dias. Renovar +10 só na renovação (o primeiro ciclo continua 30 dias). Registrar devolução.",
            "Aba Histórico: rastreio do que já voltou.",
        ],
        impact="Sem ISBN/ficha: o Cantinho fica vazio. Busca curta demais: ninguém aparece. Confirmar retirada sem o diálogo da web: o empréstimo não grava. Marcar desligado pelo is_active: falso positivo (membro ativo aparece em vermelho). Sem WhatsApp no cadastro: o ícone não abre conversa. Livro não devolvido: some do acervo disponível.",
    )

    add_heading_custom(doc, "3.C  Culto e Eventos", 1)

    add_feature(
        doc,
        code="3.C.1",
        name="Programação de Eventos",
        what="Criar, editar, replicar e excluir cultos/eventos. É o coração da operação da Secretaria. Sem evento publicado, a home, o totem, as escalas e a presença não têm âncora.",
        route="/maintenance-dashboard?panel=events",
        who="Secretaria, Gestor, Super Admin (Tesoureiro pode ajustar evento passado, conforme grant).",
        access=["Menu → engrenagem → Culto e Eventos → Programação de Eventos."],
        manages="Secretaria no dia a dia. Super Admin em qualquer caso. Tesoureiro pode furar trava de data passada (`session_can_bypass_event_past_date_lock`).",
        acl="`maintenance.card.events` (view e update). Tabelas `events`, `event_registrations`. Tela de manutenção `/maintenance-dashboard`.",
        fill=[
            "Toque em novo evento.",
            "Nome (como a congregação conhece: «Culto da Família»).",
            "Data (DD/MM/AAAA) e hora.",
            "Local e endereço (senão a home mostra «Sem local informado»).",
            "Capacidade máxima, se houver limite.",
            "Marque as salas do culto (KIDS, TEENS e as salas afetivas já cadastradas).",
            "Ligue o que for verdade no domingo: Ofertas, Totem ativo, Requer quórum, Somente membros, Geofence, Publicado.",
            "Salve. Use replicar para a série semanal, conferindo cada data.",
            "Para cancelar: exclua ou despublique — despublicado some da home.",
        ],
        impact="Não publicar: ninguém vê o culto. Totem desligado: QR inútil. Quórum desligado em assembleia: lista legal incompleta. Salas desmarcadas: kids no corredor. Data/hora erradas: família no templo vazio. Evento só-membros mal marcado: visitante barrado ou membro misturado.",
    )

    add_feature(
        doc,
        code="3.C.2",
        name="Cronograma de Eventos",
        what="Visão Gantt (linha do tempo) dos eventos. Toque abre o editor da Programação.",
        route="/maintenance-dashboard?panel=events_gantt",
        who="Os mesmos da Programação de Eventos.",
        access=["Menu → engrenagem → Culto e Eventos → Cronograma de Eventos."],
        manages="Igual à Programação: a edição real é o formulário de evento.",
        acl="`maintenance.card.events_gantt`.",
        fill=["Olhe a semana/mês. Toque no bloco para corrigir. Não cadastre evento só no Gantt — use a Programação se o toque não abrir o form."],
        impact="Sem eventos, o Gantt é vazio. Sobreposição não vista aqui vira dois cultos no mesmo horário na home.",
    )

    add_feature(
        doc,
        code="3.C.3",
        name="Manutenção de Avisos",
        what="Comunicados que aparecem na home (página Avisos) e alimentam o orquestrador do culto.",
        route="/maintenance-dashboard?panel=event_orchestration  — hint: Comunicados da home",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Culto e Eventos → Manutenção de Avisos."],
        manages="Secretaria redige e dispara na hora do culto. `profile_is_event_control_admin` = Super Admin OU papel Secretaria OU update neste card.",
        acl="`maintenance.card.event_orchestration`. Tabelas `event_control`, `event_avisos`. Tela espelho `/admin/orquestrador`.",
        fill=[
            "Crie o aviso com texto que a congregação entenda em dois segundos.",
            "Vincule ao culto quando a tela pedir.",
            "Publique. Confira na home (Avisos) com um usuário membro, se puder.",
            "No horário do culto, use o Orquestrador (3.C.9) para o painel ao vivo.",
        ],
        impact="Sem aviso: a home de comunicados seca e o telão do orquestrador não tem roteiro. Aviso errado publicado: a igreja inteira lê informação falsa.",
    )

    add_feature(
        doc,
        code="3.C.4",
        name="Sala(s) - Check In",
        what="Operação do servidor de sala no culto (KIDS/TEENS e salas afetivas): quem chegou, contato, contagem.",
        route="/maintenance-dashboard?panel=sala_servidor",
        who="Secretaria, Gestor, Super Admin — e quem estiver no aparelho da sala.",
        access=["Menu → engrenagem → Culto e Eventos → Sala(s) - Check In."],
        manages="Secretaria prepara salas e evento. O servidor da sala opera a lista na hora.",
        acl="`maintenance.card.sala_servidor`. Depende de salas cadastradas + chaves habilitadas no evento.",
        fill=[
            "Selecione o evento de hoje.",
            "Escolha a sala.",
            "Confira as crianças/adolescentes vinculados. Use WhatsApp do responsável se faltar alguém.",
            "Não «invente» check-in de quem não está na atribuição — corrija a Configuração de salas.",
        ],
        impact="Evento sem salas: painel vazio. Criança sem atribuição: some da lista. Sem este painel no domingo, a sala trabalha no caderno de papel e o app fica defasado.",
    )

    add_feature(
        doc,
        code="3.C.5",
        name="Tipos de Escala",
        what="Cadastro dos tipos (Louvor, Recepção, Midia, Kids…). Capacidade e interruptores. Base para servos e programação.",
        route="/maintenance-dashboard?panel=scale_types",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Culto e Eventos → Tipos de Escala."],
        manages="Secretaria. Código do tipo deve ser estável.",
        acl="`maintenance.card.scale_types`. Update neste card libera os tipos em `profile_has_scale_type_access`.",
        fill=[
            "Novo tipo: nome visível + código curto sem espaço.",
            "Capacidade 1–50 (quantos servem por culto).",
            "Ligue os interruptores que a tela oferecer (ativo, etc.).",
            "Salve antes de cadastrar servos.",
        ],
        impact="Sem tipos: não há como cadastrar servo nem programar. Código mudado depois: grants `scale_type.*` antigos deixam de bater.",
    )

    add_feature(
        doc,
        code="3.C.6",
        name="Servos em Disponibilidade",
        what="Quem pode ser escalado em cada tipo (o banco de pessoas, não ainda a data do culto).",
        route="/maintenance-dashboard?panel=scale_volunteers",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Culto e Eventos → Servos em Disponibilidade."],
        manages="Secretaria monta a base com os nomes de cada ministério.",
        acl="`maintenance.card.scale_volunteers`.",
        fill=[
            "Escolha o tipo de escala.",
            "Busque a pessoa pelo nome (ela precisa existir no cadastro).",
            "Vincule. Tire quem não serve mais naquele ministério.",
        ],
        impact="Base vazia: a Programação de Escalas não tem nomes. Pessoa fora do cadastro: não aparece na busca — recepcione/cadastre primeiro.",
    )

    add_feature(
        doc,
        code="3.C.7",
        name="Programação de Escalas",
        what="O culto concreto: tipo + data + servo. É o que o membro vê em Menu → Escalas.",
        route="/maintenance-dashboard?panel=scales",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Culto e Eventos → Programação de Escalas."],
        manages="Secretaria fecha a escala da semana (todos os tipos ativos).",
        acl="`maintenance.card.scales`.",
        fill=[
            "Escolha o tipo e a data do culto (a data deve existir na Programação de Eventos).",
            "Atribua o servo da lista de disponibilidade.",
            "Salve. Confira se a capacidade do tipo não estourou.",
            "Exclua substituição errada; oriente o membro a pedir troca pelo app se for o caso.",
        ],
        impact="Sem programação: o membro abre Escalas e acha que está de folga. Data sem evento: a operação do domingo desconecta. Sala servidor e telão ficam sem nomes.",
    )

    add_feature(
        doc,
        code="3.C.8",
        name="Presença",
        what="Registro de presença/quórum dos eventos que exigem quórum (assembleia).",
        route="/maintenance-dashboard?panel=quorum_presence",
        who="Secretaria, Gestor, Super Admin.",
        access=["Menu → engrenagem → Culto e Eventos → Presença."],
        manages="Secretaria no culto/assembleia. O check-in familiar e o totem alimentam a lista.",
        acl="`maintenance.card.quorum_presence`. O evento precisa `requer_quorum`.",
        fill=[
            "Selecione o evento de assembleia.",
            "Confira quem o totem/agenda já confirmou.",
            "Não desmarque à mão quem o totem cadeou sem critério legal da igreja.",
            "Gere/exporte o que a tela oferecer para ata.",
        ],
        impact="Quórum desligado no evento: esta lista não sustenta ata. Famílias não marcadas: assembleia «sem gente» no papel. Relatório de assembleia (Governança) sai furado.",
    )

    add_feature(
        doc,
        code="3.C.9",
        name="Orquestrador",
        what="Painel em tela cheia da sequência de avisos/culto ao vivo (para projetor ou operador).",
        route="/admin/orquestrador  — hint: Painel de avisos em tela cheia",
        who="Secretaria e Super Admin (quem opera o projetor no culto).",
        access=["Menu → engrenagem → Culto e Eventos → Orquestrador. Use um computador ligado ao projetor."],
        manages="Quem opera o culto no minuto. O conteúdo vem da Manutenção de Avisos / event_control.",
        acl="Tela `/admin/orquestrador` + `maintenance.card.event_orchestration`. Função SQL `profile_is_event_control_admin`.",
        fill=[
            "Abra alguns minutos antes do culto.",
            "Avance os avisos na ordem combinada com a Secretaria.",
            "Não edite texto longo aqui — corrija antes na Manutenção de Avisos.",
        ],
        impact="Painel ocioso: a congregação não é guiada. Aviso desatualizado no telão: informação pública errada.",
    )

    add_heading_custom(doc, "3.D  Finanças e Inteligência", 1)
    add_para(doc, "Tesoureiro e Super Admin são o núcleo do livro caixa e do preditivo. Secretaria vê Gestão de Campanhas (grant absorvido); não vê Informações Financeiras nem Modelo Preditivo.", italic=True, color=MUTED)

    add_feature(
        doc,
        code="3.D.1",
        name="Informações Financeiras",
        what="Livro da tesouraria: lançamentos, comprovantes, RD, orçamento, CSV, atas/PDF.",
        route="/maintenance-dashboard?panel=financials  — hint: Extratos, RD e orçamento",
        who="Tesoureiro e Super Admin. Secretaria não.",
        access=["Menu → engrenagem → Finanças e Inteligência → Informações Financeiras."],
        manages="Tesoureiro. Super Admin irrestrito. O membro só envia RD pelo Perfil.",
        acl="`maintenance.card.financials`. Tabela `financials`.",
        fill=[
            "Escolha o mês.",
            "Lance conta, valor e observação. Anexe comprovante.",
            "Trate a fila de RD: emitir, vincular, conciliar.",
            "Mantenha o orçamento (versões) se a igreja trabalhar com planejado.",
            "Importe CSV só com arquivo no layout combinado.",
        ],
        impact="Mês sem lançamento: o Financeiro do membro e o preditivo esvaziam. RD sem conciliação: irmão no prejuízo e livro mentiroso. Sem orçamento: Planejado × Realizado bloqueia.",
    )

    add_feature(
        doc,
        code="3.D.2",
        name="Gestão de Campanhas",
        what="Campanhas e projetos com meta em R$, datas, Pix identificado e status (alimenta «Eu quero…» / ofertas).",
        route="/maintenance-dashboard?panel=campaigns_management",
        who="Secretaria, Pastoral, Tesoureiro, Super Admin. Gestor em view.",
        access=["Menu → engrenagem → Finanças e Inteligência → Gestão de Campanhas."],
        manages="Secretaria ou tesouraria/pastoral donos da campanha.",
        acl="`maintenance.finance.campaigns`.",
        fill=[
            "Título e descrição que o membro entenda.",
            "Meta em reais, início e fim.",
            "Configure o Pix em centavos/identificador conforme a tela.",
            "Status ativo só quando puder receber. Encerre no fim.",
        ],
        impact="Sem campanha ativa: o atalho da home avisa que não há projeto. Meta/data vazias: contribuição sem rumo e sem relatório.",
    )

    add_feature(
        doc,
        code="3.D.3",
        name="Modelo Preditivo",
        what="Leitura de tendências (financeiro, presença, transferências). Não é formulário de lançamento.",
        route="/maintenance-dashboard?panel=predictive_insights",
        who="Tesoureiro, Pastoral, Gestor, Super Admin. Secretaria não.",
        access=["Menu → engrenagem → Finanças e Inteligência → Modelo Preditivo."],
        manages="Ninguém «preenche» o modelo: ele depende da qualidade dos lançamentos e da presença.",
        acl="`maintenance.card.predictive_insights`.",
        fill=["Abra o painel depois de fechar o mês. Use para decisão; não invente número aqui."],
        impact="Histórico fraco: painel vazio ou enganoso. Decidir obra/contrato só com preditivo oco é risco.",
    )

    add_heading_custom(doc, "3.E  Governança e TI (ordem da tela)", 1)
    add_para(
        doc,
        "Na interface, ao abrir o grupo Governança e TI, a ordem visível é: Relatórios → Controle de Acesso → Mudança Papéis → Transferência de Membro → Acesso Usuários → Modo Ghost → o agrupador «Manutenção da Trilha» (Temas, Reconhecimentos, Reset) → Assinaturas. «Instâncias (Igrejas)» fica fixo no rodapé do painel Configurações.",
        italic=True,
        color=MUTED,
    )

    add_feature(
        doc,
        code="3.E.1",
        name="Relatórios",
        what="Catálogo de extrações: membros ativos/inativos, faixa etária, necessidades pastorais, saúde infantil (LGPD), quórum, estacionamento, sugestões, inscrições em eventos.",
        route="/maintenance-dashboard?panel=relatorios",
        who="Secretaria (view), Gestor, Pastoral, Tesoureiro, Super Admin.",
        access=["Menu → engrenagem → Governança e TI → Relatórios."],
        manages="Quem tem o grant gera. A qualidade depende dos cadastros (Secretaria) e da tesouraria/pastoral.",
        acl="`maintenance.card.relatorios`.",
        fill=[
            "Escolha o relatório.",
            "Preencha parâmetros (mês, evento, recorte).",
            "Gere. Trate dado de saúde/criança como sigilo.",
        ],
        impact="Cadastro podre: PDF inútil. Quórum sem flag no evento: assembleia sem lista. Saúde infantil nas mãos erradas: incidente LGPD.",
    )

    add_feature(
        doc,
        code="3.E.2",
        name="Controle de Acesso",
        what="Matriz papéis × telas/tabelas/colunas (bolinhas de view/update) e usuários. É onde se atribui o papel Secretaria, Tesoureiro, etc.",
        route="/maintenance-dashboard?panel=access_control",
        who="Gestor em Controle de Acesso e Super Admin. Secretaria não. Gestor nunca vê Super Administrador nem PIN.",
        access=["Menu → engrenagem → Governança e TI → Controle de Acesso."],
        manages="Gestor (no recorte permitido) e Super Admin. Comentário de proteção no código: Gestor não tem visibilidade do Super Administrador.",
        acl="`maintenance.card.access_control` + RPCs `assert_gestor_super_admin_shield` / `profile_visible_to_access_actor`.",
        fill=[
            "Aba Papéis: escolha o papel, ligue só o que aquele ministério precisa (fail-closed: o que não liga, some do menu).",
            "Aba usuários: busque a pessoa e associe o papel. Uma pessoa pode somar papéis (ex.: Membro + Secretaria).",
            "Não ligue Cuidado Pastoral para Secretaria. Não ligue curinga `*` fora do Super Admin.",
            "Teste saindo e entrando (ou Ghost, se Super Admin) com o usuário alvo.",
        ],
        impact="Grant a mais: vazamento (finanças, sigilo pastoral). Grant a menos: a Secretaria não vê Eventos e o domingo para. Mexer em Super Admin pelo Gestor é bloqueado de propósito.",
    )

    add_feature(
        doc,
        code="3.E.3",
        name="Mudança Papéis",
        what="Fluxo pastoral para alterar papel/membership de um perfil (visitante → congregado → membro, etc.).",
        route="/maintenance-dashboard?panel=mudanca_papeis",
        who="Pastoral, Gestor, Super Admin. Secretaria não. Papéis protegidos contra esta tela: `super_admin`, `pastoral`, `secretaria`.",
        access=["Menu → engrenagem → Governança e TI → Mudança Papéis."],
        manages="Equipe Pastoral. Super Admin em exceção.",
        acl="`maintenance.card.mudanca_papeis`. Função `profile_has_protected_role_for_pastoral_change`.",
        fill=[
            "Busque a pessoa.",
            "Escolha o novo papel com critério (não pule visita pastoral).",
            "Confirme datas de membership se a tela pedir.",
            "Não use isto para «dar Secretaria» — isso é Controle de Acesso, e Secretaria está protegida aqui.",
        ],
        impact="Não operar: a pessoa fica eternamente visitante no ACL. Operar cedo demais: membro sem vínculo real. Tentar alterar Super Admin/Pastoral/Secretaria: a função bloqueia.",
    )

    add_feature(
        doc,
        code="3.E.4",
        name="Transferência de Membro",
        what="Pedido de transferência entre instâncias (igrejas) do ecossistema, com motivo e aceite origem/destino.",
        route="/maintenance-dashboard?panel=transferencia_igreja",
        who="Pastoral e Super Admin.",
        access=["Menu → engrenagem → Governança e TI → Transferência de Membro."],
        manages="Pastoral das duas pontas. Super Admin desempata.",
        acl="`maintenance.card.transferencia_igreja`. Isolamento por tenant até o aceite.",
        fill=[
            "Informe o celular, localize o perfil.",
            "Preencha o motivo.",
            "Envie. Acompanhe origem/destino. Cancele se foi engano.",
        ],
        impact="Sem transferência: a pessoa some numa igreja e «não existe» na outra, ou fica duplicada. Motivo vazio: a outra ponta não decide.",
    )

    add_feature(
        doc,
        code="3.E.5",
        name="Acesso Usuários",
        what="Insights de uso (quem acessou o quê) para auditoria operacional da ACL.",
        route="/maintenance-dashboard?panel=profile_access_insights",
        who="Gestor e Super Admin.",
        access=["Menu → engrenagem → Governança e TI → Acesso Usuários."],
        manages="Somente consulta/auditoria. Não é tela de «dar permissão» (isso é Controle de Acesso).",
        acl="`maintenance.card.profile_access_insights`. Gestor não vê trilhas do Super Admin.",
        fill=["Busque o usuário. Leia o histórico. Cruze com um grant estranho na matriz."],
        impact="Sem olhar esta tela, grants vazados demoram para ser descobertos.",
    )

    add_feature(
        doc,
        code="3.E.6",
        name="Modo Ghost",
        what="O Super Administrador assume a identidade efetiva de outro perfil para ver menus e dados como aquela pessoa (auditoria). Toda tela em Ghost usa o alvo, não o operador.",
        route="/maintenance-dashboard?panel=auditor",
        who="Somente Super Admin (RPC canOperateGhostMode). Secretaria, Gestor e demais não.",
        access=["Menu → engrenagem → Governança e TI → Modo Ghost."],
        manages="Só Super Admin inicia e encerra. Encerrar Ghost volta à identidade real.",
        acl="`maintenance.card.auditor`. Cliente: `loadEffectiveSessionProfile` / `resolveEffectiveProfileId` / `getEffectiveUserPhone`.",
        fill=[
            "Escolha o perfil alvo.",
            "Navegue como se fosse a pessoa (menu, home, engrenagem).",
            "Encerre o Ghost ao terminar — não lance evento «sendo» o membro.",
            "Não use Ghost para alterar PIN ou para operar tesouraria no nome alheio.",
        ],
        impact="Ghost aberto por engano: você grava presença/avisos como outra pessoa. Não encerrar: o próximo toque ainda é o alvo. Gestor sem Ghost é proteção, não falta de feature.",
    )

    add_feature(
        doc,
        code="3.E.7",
        name="Manutenção da Trilha — Temas da Trilha",
        what="Conteúdo das lições (título, texto, vídeo, ativo). Na tela, fica dentro do agrupador «Manutenção da Trilha» (toque para expandir).",
        route="/maintenance-dashboard?panel=discipleship_themes  — hint: Textos, vídeos e reflexões dos passos",
        who="Secretaria, Pastoral e Super Admin. Reset NÃO é este item.",
        access=[
            "Menu → engrenagem → Governança e TI.",
            "Toque em «Manutenção da Trilha».",
            "Toque em Temas da Trilha.",
        ],
        manages="Secretaria ou pastoral/discipulado. Super Admin em última instância.",
        acl="`maintenance.card.discipleship_themes`. Drawer: Super Admin OU canAccessPastoralCare OU grant do painel.",
        fill=[
            "Abra o passo/lição.",
            "Preencha título, texto, URL do vídeo. Marque ativo só quando o conteúdo estiver pronto.",
            "A Lição do Perfil Ministerial (5.1) precisa existir e estar clara — o mural de vagas depende dela.",
        ],
        impact="Temas vazios: a Trilha do membro é um corredor sem portas. Vídeo quebrado: a turma trava. 5.1 inexistente: Mural de Oportunidades bloqueia a igreja inteira.",
    )

    add_feature(
        doc,
        code="3.E.8",
        name="Manutenção da Trilha — Reconhecimentos",
        what="Alunos com 100% prontos para certificado (status Novo / Visto / Fechado).",
        route="/maintenance-dashboard?panel=discipleship_alerts  — hint: Alunos 100% prontos para certificado",
        who="Os mesmos dos Temas.",
        access=["Engrenagem → Governança e TI → Manutenção da Trilha → Trilha — Reconhecimentos."],
        manages="Secretaria ou pastoral/discipulado entrega o reconhecimento.",
        acl="`maintenance.card.discipleship_alerts`.",
        fill=["Abra os Novos. Confira o aluno. Marque Visto quando olhar. Fechado quando o certificado/cerimônia ocorrer."],
        impact="Fila ignorada: gente formada sem festa e sem registro. Fechar cedo demais: certificado fantasma.",
    )

    add_feature(
        doc,
        code="3.E.9",
        name="Manutenção da Trilha — Resetar Trilha",
        what="Zera o progresso de um usuário na trilha desta igreja. Ação destrutiva e rara.",
        route="/maintenance-dashboard?panel=discipleship_reset  — hint: Reiniciar progresso de um usuário nesta igreja",
        who="Somente Super Admin (drawer e SQL).",
        access=["Engrenagem → Governança e TI → Manutenção da Trilha → Resetar Trilha."],
        manages="Apenas Super Administrador.",
        acl="`maintenance.card.discipleship_reset` + filtro `isSuperAdmin` no drawer.",
        fill=[
            "Busque nome ou telefone.",
            "Confirme que é a pessoa certa e a igreja certa (tenant).",
            "Resete só em caso de perfil duplicado/teste, nunca por «preguiça de lição».",
        ],
        impact="Reset errado: apaga meses de discipulado. Sem reset quando o perfil é lixo de teste: a trilha mente nos reconhecimentos.",
    )

    add_feature(
        doc,
        code="3.E.10",
        name="Assinaturas",
        what="Planos e cobrança da igreja (Stripe): checkout, sincronização, capacidade de usuários.",
        route="/billing  — hint: Planos e cobrança da igreja",
        who="Somente Super Administrador. Tesoureiro e Secretaria não veem o item.",
        access=["Menu → engrenagem → Governança e TI → Assinaturas (depois do grupo da Trilha)."],
        manages="Super Admin / operação comercial do Conecta+.",
        acl="Drawer: `menu_billing` só se `isSuperAdmin`. Sem resource_key de painel de manutenção.",
        fill=["Escolha o plano, conclua o checkout, espere a sincronização. Não misture com o livro caixa da tesouraria local (são camadas diferentes)."],
        impact="Sem assinatura ativa: a instância pode perder capacidade/serviço. Plano errado: a igreja paga o que não usa ou estoura limite de usuários.",
    )

    add_feature(
        doc,
        code="3.E.11",
        name="Aliança Conecta Reino",
        what="Demonstrativo das assinaturas Conecta+ (cartão, baixa imediata) e do passivo de 40% às igrejas mães. A quitação da oferta de apoio ministerial é manual, em até 30 dias.",
        route="/alianca-conecta-reino  — hint: Indicações, passivo de 40% e baixa manual das ofertas",
        who="Somente Super Administrador. Tesoureiro e Secretaria não veem o item. A igreja mãe lê o recorte no boletim Financeiro (seção Aliança), sem baixar oferta.",
        access=["Menu → engrenagem → Governança e TI → Aliança Conecta Reino (junto de Assinaturas, depois da Trilha)."],
        manages="Super Admin efetiva a oferta («Marcar como paga»). O Stripe baixa a assinatura do plano automaticamente; o passivo de 40% não se paga sozinho.",
        acl="Drawer `menu_alianca` só `isSuperAdmin`. RPCs `get_alianca_admin_statement`, `settle_alianca_payout_admin`. Webhook `process_alianca_invoice_paid` / `_failed`.",
        fill=[
            "Abra o demonstrativo: assinaturas x passivo.",
            "Para quitar: confirme «Marcar como paga» (o diálogo vale também no navegador).",
            "Não misture com Informações Financeiras da tesouraria local — são camadas diferentes.",
            "O ciclo da parceria encerra no 4º pagamento conforme a regra da tela.",
        ],
        impact="Sem baixa manual: o passivo de 40% envelhece e a igreja mãe não recebe. Baixar a oferta errada: ciclo da parceria avança indevido. Sem Super Admin, a operação comercial da rede some do menu de propósito.",
    )

    add_feature(
        doc,
        code="3.E.12",
        name="Instâncias (Igrejas)",
        what="Criar e alternar ambientes de igreja: logo, CNPJ/Pix, redes, ativar/inativar tenant. Na interface fica fixo no rodapé de Configurações, fora dos acordeões.",
        route="/igrejas  — hint: Criar e alternar ambientes de igreja",
        who="Somente Super Administrador.",
        access=["Menu → engrenagem → (rodapé do painel Configurações) Instâncias (Igrejas)."],
        manages="Super Admin. Cada igreja isolada por `tenant_id`.",
        acl="`useIgrejasAdminAccess` / `menu_igrejas` → só `isSuperAdmin`.",
        fill=[
            "Crie a igreja com nome oficial.",
            "Ative o tenant.",
            "Envie o logo (aparece no topo da home).",
            "Cadastre Pix (ofertas), site, Instagram, YouTube.",
            "Não opere dados de uma igreja logado em outra.",
        ],
        impact="Sem instância ativa: não há Conecta+ daquela igreja. Sem Pix: ofertas falham. Sem redes: Menu → Redes Sociais fica vazio. Logo ausente: chrome sem marca.",
    )
    doc.add_page_break()


def add_closing(doc: Document) -> None:
    add_heading_custom(doc, "4. Roteiro rápido da Secretaria no domingo", 1)
    add_para(doc, "Checklist didático — faça nesta ordem se você for Secretaria:")
    add_steps(
        doc,
        [
            "Na sexta: Programação de Eventos — culto publicado, horário, local, salas, Totem ativo, quórum se for assembleia.",
            "Programação de Escalas daquele culto (servos já cadastrados em Disponibilidade).",
            "Manutenção de Avisos do boletim da home / telão. No culto, abra o Orquestrador no projetor.",
            "No sábado: Recepção Familiar — zere a fila de lotes. Corrija CEP no Cadastro de Usuário.",
            "No hall: Totem aberto, câmera ok. Famílias com Carteirinha.",
            "Nas salas: Sala(s) - Check In no evento de hoje.",
            "Depois do culto: Presença se houver quórum. Régua de Acolhimento dos visitantes novos. Empréstimos de livros: confirmar retiradas e WhatsApp de prazo.",
            "Nunca abra Cuidados Pastorais, Informações Financeiras, Ghost, Aliança ou Assinaturas — não é o seu papel e o menu não mostra.",
        ],
    )

    add_heading_custom(doc, "5. O que este manual deliberadamente não ensina a «ligar de novo»", 1)
    add_para(
        doc,
        "O carrossel antigo do Painel (/(tabs)/dashboard) não é produto publicado: a rota só redireciona. Cards congelados (QR avulso no dashboard, Salas no carrossel, Estacionamento como card, escala avulsa duplicada) não devem ser reativados sem ordem expressa. Check-in por QR vive na Carteirinha; estacionamento vive dentro de Escalas quando o tipo pede placa.",
    )

    add_heading_custom(doc, "6. Fontes no código (para quem for atualizar este manual)", 1)
    sources = [
        "lib/appDrawerMenu.ts — ordem do Menu e da Engrenagem, rotas, panel=.",
        "components/minimal/AppDrawer.tsx e AppDrawerSettings.tsx — engrenagem, grupos, Trilha agrupada, Instâncias no rodapé.",
        "hooks/useAppDrawerMenu.ts e lib/drawerMenuAccess.ts — quem vê o quê.",
        "lib/screenAccessResourceKeys.ts e lib/accessScreen.ts — chaves ACL.",
        "lib/accessRoleDisplayOrder.ts e scripts/access-control-*.sql — papéis atuais (Secretaria absorveu Líder / Líder Geral / Administrador de Eventos / Orquestrador).",
        "app/(tabs)/index.tsx, EventsInboxHome, MinimalEuQueroFooter — home fora do menu.",
        "app/maintenance-dashboard.tsx — painéis da engrenagem.",
        "app/sobre-conecta.tsx e lib/conectaPrivacyDeclaration.ts — declaração LGPD de consulta.",
        "app/livros-doados.tsx, components/LivrosDoadosPanel.tsx, LivrosEmprestimosPanel.tsx, MeusLivrosRetiradosPanel.tsx — acervo e empréstimos.",
        "lib/isbnCatalogLookup.ts e functions/api/buscar-livro.ts — Google Books, CBL/BrasilAPI, Open Library.",
        "app/alianca-conecta-reino.tsx e lib/alianca/ — passivo 40% e baixa manual.",
    ]
    for s in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r = p.add_run(s)
        set_run_font(r, size=11)

    add_para(
        doc,
        "Fim do guia. Se a tela do seu usuário estiver diferente, a causa mais comum é o papel (ACL) ou o vínculo de membro inativo — não um «bug de menu».",
        space_before=16,
        italic=True,
        color=MUTED,
    )


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_howto(doc)
    add_part0(doc)
    add_part1(doc)
    add_part2(doc)
    add_part3(doc)
    add_closing(doc)
    doc.save(str(OUT_PATH))
    print(f"Salvo em: {OUT_PATH}")


if __name__ == "__main__":
    main()
