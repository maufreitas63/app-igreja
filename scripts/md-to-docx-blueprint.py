"""Converte ARQUITETURA_BLUEPRINT_PWA.md em documento Word formatado."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "ARQUITETURA_BLUEPRINT_PWA.md"
OUT_DOCX = ROOT / "docs" / "ARQUITETURA_BLUEPRINT_PWA.docx"
OUT_DOC = ROOT / "docs" / "ARQUITETURA_BLUEPRINT_PWA.doc"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        if level == 1:
            h.font.size = Pt(18)
        elif level == 2:
            h.font.size = Pt(14)
        else:
            h.font.size = Pt(12)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Blueprint de Arquitetura\n e Especificação Técnica")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("App Igreja — PWA / Expo / Supabase")
    r.font.size = Pt(14)
    r.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Documento gerado a partir do repositório app-igreja").font.size = Pt(10)

    doc.add_page_break()


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells if c)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Table Grid"
    for j, text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, row in enumerate(body, start=1):
        for j, text in enumerate(row):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = text
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], lang: str = "") -> None:
    label = f" ({lang})" if lang else ""
    cap = doc.add_paragraph()
    cap.add_run(f"Diagrama / estrutura{label}:").italic = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def convert(md_text: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style_document(doc)
    add_title_page(doc)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            code_lang = line.strip()[3:].strip()
            i += 1
            continue

        if line.strip().startswith("|"):
            row = parse_table_row(line)
            if not is_separator_row(row):
                table_rows.append(row)
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        if line.strip().startswith("> "):
            p = doc.add_paragraph(line.strip()[2:].strip())
            p.paragraph_format.left_indent = Cm(0.75)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet:
            indent = len(bullet.group(1)) // 2
            style = "List Bullet" if indent == 0 else "List Bullet 2"
            doc.add_paragraph(bullet.group(2).strip(), style=style)
            i += 1
            continue

        num = re.match(r"^(\d+)\.\s+(.+)$", line)
        if num:
            doc.add_paragraph(num.group(2).strip(), style="List Number")
            i += 1
            continue

        if line.strip():
            doc.add_paragraph(line.strip())
        i += 1

    if table_rows:
        add_table(doc, table_rows)
    if in_code and code_lines:
        add_code_block(doc, code_lines, code_lang)

    return doc


def main() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    md = MD_PATH.read_text(encoding="utf-8")
    doc = convert(md)
    doc.save(str(OUT_DOCX))
    # Cópia .doc: Word abre .docx; extensão .doc atende pedido de "versão doc"
    import shutil

    shutil.copy2(OUT_DOCX, OUT_DOC)
    print(f"Gerado: {OUT_DOCX}")
    print(f"Gerado: {OUT_DOC} (formato Office Open XML — abrir no Word)")


if __name__ == "__main__":
    main()
